"""Build script for classic-zh docxtpl template.

Run once to produce template.docx. Keep in repo — it's the source of truth for
regenerating the template.

Usage:
    cd assets/templates/classic-zh
    python _build_template.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

HERE = Path(__file__).parent
OUT = HERE / "template.docx"


def set_cjk_font(run, font_name: str = "宋体") -> None:
    """Set CJK (eastAsia) font and also ascii/hAnsi to Times New Roman."""
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def add_styled_paragraph(doc, text: str, *, size: float, bold: bool = False,
                          align=WD_PARAGRAPH_ALIGNMENT.LEFT,
                          space_before: float = 0, space_after: float = 2) -> None:
    """Add a paragraph with a single run, styled via python-docx."""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(size * 1.3)

    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    set_cjk_font(run)


def add_section_heading(doc, title: str) -> None:
    """Section heading: 12pt bold, with a bottom border line via XML."""
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(16)

    run = p.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    set_cjk_font(run)

    # Add bottom border to visually separate sections
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_body_paragraph(doc, text: str, *, indent: bool = False,
                        space_before: float = 0, space_after: float = 1) -> None:
    """Body text at 10.5pt."""
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(14)
    if indent:
        pf.left_indent = Cm(0.5)

    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    set_cjk_font(run)


def _add_exp_header(doc: Document) -> None:
    """Experience header: title left, time right-aligned via tab stop at 18cm."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(1)
    pf.line_spacing = Pt(14)
    # Right-aligned tab stop at page usable width: 21cm - 2×1.5cm = 18cm
    pf.tab_stops.add_tab_stop(Cm(18), WD_TAB_ALIGNMENT.RIGHT)

    run_title = p.add_run("{{ e.title }}")
    run_title.bold = True
    run_title.font.name = "Times New Roman"
    run_title.font.size = Pt(10.5)
    set_cjk_font(run_title)

    run_tab = p.add_run("\t")
    run_tab.font.size = Pt(10.5)

    run_time = p.add_run("{{ e.fields.时间 }}")
    run_time.font.name = "Times New Roman"
    run_time.font.size = Pt(10.5)
    set_cjk_font(run_time)


def build() -> None:
    doc = Document()

    # ── Page setup: A4, 1.5cm margins ─────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Remove default empty paragraph added by python-docx
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # ── 1. Name ────────────────────────────────────────────────────────────
    add_styled_paragraph(
        doc,
        "{{ profile.basic.name }}",
        size=22,
        bold=True,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        space_before=0,
        space_after=4,
    )

    # ── 2. Basic info bar ──────────────────────────────────────────────────
    add_styled_paragraph(
        doc,
        "{{ profile.basic.gender }} | {{ profile.basic.birthdate }} | {{ profile.basic.phone }} | {{ profile.basic.email }}",
        size=10.5,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        space_before=0,
        space_after=4,
    )

    # ── 3. Education ───────────────────────────────────────────────────────
    add_section_heading(doc, "教育背景")

    add_body_paragraph(
        doc,
        "{{ profile.academics.school }}（{{ profile.academics.college }}） {{ profile.academics.major }}    {{ profile.academics.entry_date }}",
    )
    add_body_paragraph(
        doc,
        "GPA：{{ profile.academics.gpa }}   排名：{{ profile.academics.rank }}",
        space_after=2,
    )

    # ── 4. Core experiences ────────────────────────────────────────────────
    add_section_heading(doc, "核心经历")

    # docxtpl paragraph-level loop: each {%p ... %} must be the sole content of its paragraph
    _tpl_para(doc, "{%p for e in profile.experiences %}")

    # Experience title + time: title left, time right-aligned via tab stop
    # Page width 21cm - left margin 1.5cm - right margin 1.5cm = 18cm usable width
    _add_exp_header(doc)

    # Role line
    add_body_paragraph(doc, "角色：{{ e.fields.角色 }}")

    # Bullets inner loop
    _tpl_para(doc, "{%p for b in e.bullets %}")
    add_body_paragraph(doc, "• {{ b }}", indent=False)
    _tpl_para(doc, "{%p endfor %}")

    _tpl_para(doc, "{%p endfor %}")

    # ── 5. Student work ────────────────────────────────────────────────────
    add_section_heading(doc, "学生工作")

    _tpl_para(doc, "{%p for w in profile.student_work %}")
    add_body_paragraph(doc, "• {{ w }}")
    _tpl_para(doc, "{%p endfor %}")

    # ── 6. Skills ──────────────────────────────────────────────────────────
    add_section_heading(doc, "技能")

    _tpl_para(doc, "{%p for s in profile.skills %}")
    add_body_paragraph(doc, "• {{ s }}")
    _tpl_para(doc, "{%p endfor %}")

    # ── 7. Awards ──────────────────────────────────────────────────────────
    add_section_heading(doc, "奖项与荣誉")

    _tpl_para(doc, "{%p for a in profile.awards %}")
    add_body_paragraph(doc, "• {{ a }}")
    _tpl_para(doc, "{%p endfor %}")

    doc.save(str(OUT))
    print(f"Template written to {OUT}")


def _tpl_para(doc: Document, tag: str) -> None:
    """Add a paragraph containing ONLY a docxtpl block tag — no extra whitespace."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    run = p.add_run(tag)
    run.font.size = Pt(1)   # near-invisible; docxtpl replaces/removes the paragraph


if __name__ == "__main__":
    build()
